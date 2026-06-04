import re
import json
from io import BytesIO
from docx import Document
import PyPDF2
from PIL import Image
import numpy as np

def extract_key_value(text):
    data = {}
    patterns = {
        "age": r"age[:\-]?\s*(\d+)",
        "gender": r"gender[:\-]?\s*(male|female)",
        "hypertension": r"hypertension[:\-]?\s*(\d+)",
        "heart_disease": r"heart\s*disease[:\-]?\s*(\d+)",
        "smoking_history": r"smoking\s*history[:\-]?\s*([\w\s]+)",
        "bmi": r"bmi[:\-]?\s*([\d.]+)",
        "blood_glucose_level": r"blood\s*glucose\s*level[:\-]?\s*([\d.]+)"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(1).strip()
            if key in ["age", "hypertension", "heart_disease"]:
                data[key] = int(value)
            elif key in ["bmi", "blood_glucose_level"]:
                data[key] = float(value)
            else:
                data[key] = value

    return {"data": data}

def convert_to_json(file_bytes, file_type):
    text = ""

    if file_type == "txt":
        text = file_bytes.decode("utf-8")

    elif file_type == "pdf":
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            text += page.extract_text()

    elif file_type == "doc":
        doc = Document(BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"

    return json.dumps(extract_key_value(text)).encode("utf-8")

def extract_text(file_bytes, file_type):
    text = ""

    if file_type == "txt":
        text = file_bytes.decode("utf-8", errors="ignore")

    elif file_type == "pdf":
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            text += page.extract_text() or ""

    elif file_type in ["doc", "docx"]:
        doc = Document(BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"

    return text.strip()

def preprocess_image(file_bytes, target_size=(224, 224)):
    """
    Preprocess image for AI models:
    - Convert to RGB
    - Resize to 224x224 (default)
    - Normalize pixel values 0–1
    """
    image = Image.open(BytesIO(file_bytes)).convert("RGB")
    image = image.resize(target_size)
    img_array = np.array(image) / 255.0
    return img_array